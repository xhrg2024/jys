"""
评估问答系统效果（含完整 Planner 调用日志）。
用法：python src/evaluation/evaluate_qa.py

输出：
  - data/eval_results.json    — 详细 JSON 结果
  - data/eval_report.docx     — Word 报告（需 python-docx）
  - data/eval_report.md       — Markdown 报告（兜底）
"""
import json
import os
import requests
import time
from datetime import datetime
from pathlib import Path
from dotenv import load_dotenv

# 加载 .env 文件（从项目根目录）
dotenv_path = Path(__file__).resolve().parents[2] / ".env"
load_dotenv(dotenv_path)

API_PORT = int(os.environ.get("API_PORT", 8000))
API_URL = f"http://localhost:{API_PORT}/chat"

# ══════════════════════════════════════════════
# 测试集：30 题，覆盖五种意图 + SQL/图/向量检索
# ══════════════════════════════════════════════
TEST_CASES = [
    # ═══ 事实问答 (FACTUAL) — 10 题 ═══
    # 人物
    {"id": 1,  "type": "FACTUAL", "question": "马国翰是谁？",
     "keywords": ["清代", "辑佚", "玉函山房辑佚书"], "expect_sql": True},
    {"id": 2,  "type": "FACTUAL", "question": "王应麟是谁？",
     "keywords": ["宋代", "辑佚", "鼻祖"], "expect_sql": True},
    {"id": 3,  "type": "FACTUAL", "question": "严可均的代表作是什么？",
     "keywords": ["全上古", "秦汉三国六朝文"], "expect_sql": True},
    {"id": 4,  "type": "FACTUAL", "question": "王仁俊在辑佚方面有什么贡献？",
     "keywords": ["玉函山房", "续编", "补编"], "expect_sql": True},
    # 文献
    {"id": 5,  "type": "FACTUAL", "question": "《永乐大典》是什么类型的文献？",
     "keywords": ["类书", "明代"], "expect_sql": True},
    {"id": 6,  "type": "FACTUAL", "question": "《玉函山房辑佚书》的作者是谁？",
     "keywords": ["马国翰"], "expect_sql": True},
    {"id": 7,  "type": "FACTUAL", "question": "《敬齋古今黈》的作者是谁？",
     "keywords": ["敬齋"], "expect_sql": True},
    {"id": 8,  "type": "FACTUAL", "question": "《鹤林玉露》是什么朝代的文献？",
     "keywords": ["鹤林", "明代"], "expect_sql": True},
    {"id": 9,  "type": "FACTUAL", "question": "《全上古三代秦汉三国六朝文》的编纂者是谁？",
     "keywords": ["严可均", "全上古"], "expect_sql": True},
    {"id": 10, "type": "FACTUAL", "question": "《芥隱筆記》属于什么类型的著作？",
     "keywords": ["笔记", "芥隱"], "expect_sql": True},

    # ═══ 关系问答 (RELATION) — 5 题 ═══
    {"id": 11, "type": "RELATION", "question": "马国翰和王仁俊有什么关系？",
     "keywords": ["续编", "补编", "传承"], "expect_sql": True},
    {"id": 12, "type": "RELATION", "question": "王应麟对清代辑佚学有什么影响？",
     "keywords": ["鼻祖", "影响"], "expect_sql": False},
    {"id": 13, "type": "RELATION", "question": "《全宋文》与《永乐大典》有什么关系？",
     "keywords": ["底本", "来源", "辑出"], "expect_sql": True},
    {"id": 14, "type": "RELATION", "question": "惠栋和马国翰在辑佚方法上有什么关联？",
     "keywords": ["惠栋", "马国翰", "辑佚"], "expect_sql": True},
    {"id": 15, "type": "RELATION", "question": "章学诚与辑佚学有什么关系？",
     "keywords": ["章学诚", "校雠", "辑佚"], "expect_sql": False},

    # ═══ 脉络问答 (CHAIN) — 5 题 ═══
    {"id": 16, "type": "CHAIN", "question": "清代辑佚学经历了哪几个阶段？",
     "keywords": ["阶段", "发展", "乾嘉"], "expect_sql": True},
    {"id": 17, "type": "CHAIN", "question": "辑佚学的发展历程是怎样的？",
     "keywords": ["宋代", "清代", "发展"], "expect_sql": False},
    {"id": 18, "type": "CHAIN", "question": "明代类书编纂经历了怎样的发展？",
     "keywords": ["永乐", "明代", "类书"], "expect_sql": True},
    {"id": 19, "type": "CHAIN", "question": "清代辑佚学从清初到清末有什么变化？",
     "keywords": ["清初", "乾嘉", "清末"], "expect_sql": True},
    {"id": 20, "type": "CHAIN", "question": "宋代的类书编纂脉络是怎样的？",
     "keywords": ["宋代", "类书", "编纂"], "expect_sql": True},

    # ═══ 方法问答 (METHOD) — 5 题 ═══
    {"id": 21, "type": "METHOD", "question": "辑佚的基本方法有哪些？",
     "keywords": ["辑佚", "方法"], "expect_sql": False},
    {"id": 22, "type": "METHOD", "question": "校勘在辑佚中起什么作用？",
     "keywords": ["校理", "纠正", "校勘"], "expect_sql": False},
    {"id": 23, "type": "METHOD", "question": "什么是辑佚的「非彼即此推论法」？",
     "keywords": ["推论", "辑佚"], "expect_sql": False},
    {"id": 24, "type": "METHOD", "question": "清代辑佚有哪些常用的方法和途径？",
     "keywords": ["辑佚", "方法", "清代"], "expect_sql": False},
    {"id": 25, "type": "METHOD", "question": "如何进行类书辑佚？",
     "keywords": ["类书", "辑佚", "方法"], "expect_sql": True},

    # ═══ 比较问答 (COMPARE) — 5 题 ═══
    {"id": 26, "type": "COMPARE", "question": "马国翰和严可均在辑佚方面有什么不同？",
     "keywords": ["不同", "区别"], "expect_sql": True},
    {"id": 27, "type": "COMPARE", "question": "明代辑佚和清代辑佚有什么差异？",
     "keywords": ["差异", "区别"], "expect_sql": True},
    {"id": 28, "type": "COMPARE", "question": "惠栋和王谟的辑佚方法有什么异同？",
     "keywords": ["惠栋", "王谟", "辑佚"], "expect_sql": True},
    {"id": 29, "type": "COMPARE", "question": "类书辑佚和古注辑佚有什么不同？",
     "keywords": ["类书", "古注", "辑佚"], "expect_sql": False},
    {"id": 30, "type": "COMPARE", "question": "宋代的类书和明代的类书有什么差异？",
     "keywords": ["宋代", "明代", "类书"], "expect_sql": True},
]


def call_api(question):
    """调用问答API，返回 (answer, plan_log, elapsed)"""
    try:
        start = time.time()
        resp = requests.post(API_URL, json={"question": question}, timeout=120)
        elapsed = time.time() - start
        if resp.status_code == 200:
            data = resp.json()
            return data.get("answer", ""), data.get("plan_log", {}), elapsed
        else:
            return f"ERROR: HTTP {resp.status_code}", {}, elapsed
    except Exception as e:
        return f"ERROR: {str(e)}", {}, 0


def evaluate_answer(answer, keywords):
    """评估答案质量"""
    score = {
        "has_keyword": 0,
        "keyword_count": 0,
        "keyword_hits": [],
        "has_source": 0,
        "no_hallucination": 1,
        "length_ok": 0,
    }
    hit = 0
    for kw in keywords:
        if kw in answer:
            hit += 1
            score["keyword_hits"].append(kw)
    score["keyword_count"] = hit
    score["has_keyword"] = hit / len(keywords) if keywords else 0
    if "来源" in answer or "据记载" in answer or "据考证" in answer:
        score["has_source"] = 1
    for hw in ["魏源", "马国瀚", "马国瑾", "王应鳞"]:
        if hw in answer:
            score["no_hallucination"] = 0
            break
    if 50 < len(answer) < 2000:
        score["length_ok"] = 1
    return score


def format_plan_log_text(plan_log):
    """Planner 日志 → 纯文本"""
    if not plan_log:
        return "(无 Plan 日志)"
    lines = []
    lines.append(f"意图: {plan_log.get('intent', '?')}")
    lines.append(f"匹配实体: {plan_log.get('entities_matched', [])}")
    lines.append(f"图查询: {plan_log.get('graph_count', 0)}条  "
                 f"SQL查询: {plan_log.get('sql_count', 0)}条  "
                 f"向量检索: {plan_log.get('vector_count', 0)}条")
    for d in plan_log.get("sql_details", []):
        lines.append(f"  [SQL] {d[:300]}")
    for d in plan_log.get("graph_details", []):
        lines.append(f"  [图] {d[:200]}")
    return "\n".join(lines)


# ══════════════════════════════════════════════
# DOCX 报告生成
# ══════════════════════════════════════════════

def generate_docx(results, type_stats, total_stats, output_path):
    """生成 Word 评估报告"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
    except ImportError:
        return False

    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    # ── 标题 ──
    title = doc.add_heading("辑佚史智能体 — 问答系统评估报告", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"评估时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  测试集: {total_stats['total']} 题").font.size = Pt(10)

    doc.add_paragraph()

    # ── 总体指标 ──
    doc.add_heading("一、总体指标", level=1)
    table = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    metrics = [
        ("测试题目总数", str(total_stats["total"])),
        ("平均关键词命中率", f"{total_stats['avg_keyword']:.0%}"),
        ("来源标注率", f"{total_stats['source_rate']:.0%}"),
        ("无幻觉率", f"{total_stats['hallu_rate']:.0%}"),
        ("SQL 命中率", f"{total_stats['sql_rate']:.0%}"),
        ("平均响应时间", f"{total_stats['avg_time']:.1f}s"),
        ("API 可用性", f"{total_stats['api_ok']}/{total_stats['total']}"),
    ]
    for i, (label, value) in enumerate(metrics):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
        for cell in [table.cell(i, 0), table.cell(i, 1)]:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()

    # ── 按意图统计 ──
    doc.add_heading("二、按意图类型统计", level=1)
    t_table = doc.add_table(rows=len(type_stats) + 1, cols=7, style="Light Grid Accent 1")
    t_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["意图", "数量", "关键词命中率", "来源标注率", "无幻觉率", "SQL命中率", "平均耗时"]
    for j, h in enumerate(headers):
        t_table.cell(0, j).text = h
        for p in t_table.cell(0, j).paragraphs:
            for r in p.runs:
                r.font.size = Pt(9)
                r.bold = True
    for i, (t, s) in enumerate(sorted(type_stats.items()), 1):
        avg_kw = sum(s["keyword_scores"]) / len(s["keyword_scores"])
        source_rate = s["source_count"] / s["count"]
        hallu_rate = s["no_hallucination_count"] / s["count"]
        sql_rate = s["sql_hit_count"] / s["count"]
        avg_t = sum(s["times"]) / len(s["times"])
        vals = [t, str(s["count"]), f"{avg_kw:.0%}", f"{source_rate:.0%}",
                f"{hallu_rate:.0%}", f"{sql_rate:.0%}", f"{avg_t:.1f}s"]
        for j, v in enumerate(vals):
            t_table.cell(i, j).text = v
            for p in t_table.cell(i, j).paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)

    doc.add_paragraph()

    # ── 逐题详情 ──
    doc.add_heading("三、逐题详情", level=1)
    for r in results:
        # 题目标题
        h = doc.add_heading(
            f"[{r['id']}] {r['type']} — {r['question']}", level=2
        )
        for run in h.runs:
            run.font.size = Pt(11)

        # 基础信息表
        info_table = doc.add_table(rows=6, cols=2, style="Light Grid Accent 1")
        info_data = [
            ("意图", r["plan_log"].get("intent", "?") if r["plan_log"] else "?"),
            ("匹配实体", str(r["plan_log"].get("entities_matched", [])[:5] if r["plan_log"] else [])),
            ("工具调用", f"图:{r['plan_log'].get('graph_count',0)}  SQL:{r['plan_log'].get('sql_count',0)}  向量:{r['plan_log'].get('vector_count',0)}" if r["plan_log"] else "无"),
            ("期望关键词", ", ".join(r["keywords_expected"])),
            ("命中关键词", ", ".join(r["keywords_hit"]) if r["keywords_hit"] else "无"),
            ("耗时 / 来源 / 幻觉", f"{r['time']:.1f}s  |  {'✓来源' if r['score']['has_source'] else '✗来源'}  |  {'✓' if r['score']['no_hallucination'] else '✗幻觉'}"),
        ]
        for i, (label, value) in enumerate(info_data):
            info_table.cell(i, 0).text = label
            info_table.cell(i, 1).text = value

        # SQL 命中详情
        if r["plan_log"] and r["plan_log"].get("sql_details"):
            doc.add_paragraph("SQL 命中详情:", style="List Bullet")
            for d in r["plan_log"]["sql_details"]:
                p = doc.add_paragraph(d[:400])
                p.paragraph_format.left_indent = Cm(1)
                for run in p.runs:
                    run.font.size = Pt(9)

        # 回答
        doc.add_paragraph("模型回答:", style="List Bullet")
        ans_para = doc.add_paragraph()
        ans_para.paragraph_format.left_indent = Cm(0.5)
        ans_text = r["answer_full"] if len(r["answer_full"]) <= 800 else r["answer_full"][:800] + "..."
        ans_para.add_run(ans_text).font.size = Pt(9)

        doc.add_paragraph()  # 间距

    # ── SQL 覆盖分析 ──
    doc.add_heading("四、SQL 覆盖分析", level=1)
    sql_expected = [r for r in results if r.get("expect_sql")]
    sql_missed = [r for r in sql_expected if not r["sql_hit"]]
    p = doc.add_paragraph(f"期望 SQL 命中: {len(sql_expected)} 题  |  实际命中: {len(sql_expected) - len(sql_missed)} 题  |  未命中: {len(sql_missed)} 题")
    if sql_missed:
        doc.add_paragraph("SQL 未命中列表:", style="List Bullet")
        for r in sql_missed:
            doc.add_paragraph(f"[{r['id']}] {r['type']}: {r['question']}", style="List Bullet")

    doc.save(output_path)
    return True


# ══════════════════════════════════════════════
# Markdown 报告生成（兜底）
# ══════════════════════════════════════════════

def generate_markdown(results, type_stats, total_stats, output_path):
    """生成 Markdown 评估报告"""
    lines = []
    lines.append("# 辑佚史智能体 — 问答系统评估报告")
    lines.append(f"\n**评估时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  **测试集**: {total_stats['total']} 题\n")

    # 总体指标
    lines.append("## 一、总体指标\n")
    lines.append("| 指标 | 值 |")
    lines.append("|------|----|")
    lines.append(f"| 测试题目总数 | {total_stats['total']} |")
    lines.append(f"| 平均关键词命中率 | {total_stats['avg_keyword']:.0%} |")
    lines.append(f"| 来源标注率 | {total_stats['source_rate']:.0%} |")
    lines.append(f"| 无幻觉率 | {total_stats['hallu_rate']:.0%} |")
    lines.append(f"| SQL 命中率 | {total_stats['sql_rate']:.0%} |")
    lines.append(f"| 平均响应时间 | {total_stats['avg_time']:.1f}s |")
    lines.append(f"| API 可用性 | {total_stats['api_ok']}/{total_stats['total']} |")

    # 按意图统计
    lines.append("\n## 二、按意图类型统计\n")
    lines.append("| 意图 | 数量 | 关键词命中率 | 来源标注率 | 无幻觉率 | SQL命中率 | 平均耗时 |")
    lines.append("|------|------|-------------|-----------|---------|----------|---------|")
    for t, s in sorted(type_stats.items()):
        avg_kw = sum(s["keyword_scores"]) / len(s["keyword_scores"])
        source_rate = s["source_count"] / s["count"]
        hallu_rate = s["no_hallucination_count"] / s["count"]
        sql_rate = s["sql_hit_count"] / s["count"]
        avg_t = sum(s["times"]) / len(s["times"])
        lines.append(f"| {t} | {s['count']} | {avg_kw:.0%} | {source_rate:.0%} | {hallu_rate:.0%} | {sql_rate:.0%} | {avg_t:.1f}s |")

    # 逐题详情
    lines.append("\n## 三、逐题详情\n")
    for r in results:
        lines.append(f"### [{r['id']}] {r['type']} — {r['question']}\n")
        pl = r["plan_log"] if r["plan_log"] else {}
        lines.append(f"- **意图**: {pl.get('intent', '?')}")
        lines.append(f"- **匹配实体**: {pl.get('entities_matched', [])[:5]}")
        lines.append(f"- **工具调用**: 图:{pl.get('graph_count',0)}  SQL:{pl.get('sql_count',0)}  向量:{pl.get('vector_count',0)}")
        lines.append(f"- **期望关键词**: {', '.join(r['keywords_expected'])}")
        lines.append(f"- **命中关键词**: {', '.join(r['keywords_hit']) if r['keywords_hit'] else '无'}")
        lines.append(f"- **评分**: {'✓来源' if r['score']['has_source'] else '✗来源'} | {'✓' if r['score']['no_hallucination'] else '✗幻觉'} | {r['time']:.1f}s")

        if pl.get("sql_details"):
            lines.append(f"\n**SQL 命中详情**:")
            for d in pl["sql_details"]:
                lines.append(f"> 🗄️ {d[:400]}")
        if pl.get("graph_details"):
            lines.append(f"\n**图查询命中详情**:")
            for d in pl["graph_details"]:
                lines.append(f"> 🔗 {d[:200]}")

        lines.append(f"\n**模型回答**:")
        ans = r["answer_full"] if len(r["answer_full"]) <= 800 else r["answer_full"][:800] + "..."
        lines.append(f"> {ans}\n")
        lines.append("---\n")

    # SQL 覆盖分析
    lines.append("## 四、SQL 覆盖分析\n")
    sql_expected = [r for r in results if r.get("expect_sql")]
    sql_missed = [r for r in sql_expected if not r["sql_hit"]]
    lines.append(f"- **期望 SQL 命中**: {len(sql_expected)} 题")
    lines.append(f"- **实际 SQL 命中**: {len(sql_expected) - len(sql_missed)} 题")
    if sql_missed:
        lines.append(f"- **SQL 未命中** ({len(sql_missed)} 题):")
        for r in sql_missed:
            lines.append(f"  - [{r['id']}] {r['type']}: {r['question']}")

    with open(output_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return True


# ══════════════════════════════════════════════
# 主流程
# ══════════════════════════════════════════════

def main():
    print("=" * 70)
    print("  问答系统评估（含 Planner 调用日志）")
    print(f"  时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"  测试集: {len(TEST_CASES)} 题")
    print("=" * 70)

    results = []
    total_time = 0
    sql_hit_cases = 0
    api_ok = 0

    for case in TEST_CASES:
        print(f"\n{'─'*70}")
        print(f"[{case['id']}/{len(TEST_CASES)}] {case['type']}: {case['question']}")
        print(f"  期望SQL: {'是' if case.get('expect_sql') else '否'}")

        answer, plan_log, elapsed = call_api(case["question"])
        total_time += elapsed

        # 每个请求间隔1秒，避免资源累积
        time.sleep(1)

        if not answer.startswith("ERROR"):
            api_ok += 1
            score = evaluate_answer(answer, case["keywords"])
        else:
            score = {"has_keyword": 0, "keyword_count": 0, "keyword_hits": [],
                     "has_source": 0, "no_hallucination": 1, "length_ok": 0}
        score["time"] = elapsed

        sql_hit = plan_log.get("sql_count", 0) > 0 if plan_log else False
        if sql_hit:
            sql_hit_cases += 1

        results.append({
            "id": case["id"],
            "type": case["type"],
            "question": case["question"],
            "answer": answer[:500],
            "answer_full": answer,
            "keywords_expected": case["keywords"],
            "keywords_hit": score["keyword_hits"],
            "score": score,
            "time": elapsed,
            "plan_log": plan_log,
            "expect_sql": case.get("expect_sql", False),
            "sql_hit": sql_hit,
        })

        # ── 终端输出 ──
        print(format_plan_log_text(plan_log))
        print(f"  回答: {answer[:150]}{'...' if len(answer) > 150 else ''}")
        print(f"  关键词命中: {score['keyword_count']}/{len(case['keywords'])} "
              f"({score['has_keyword']:.0%}) → {score['keyword_hits']}")
        print(f"  来源标注: {'✓' if score['has_source'] else '✗'}  "
              f"无幻觉: {'✓' if score['no_hallucination'] else '✗'}  "
              f"SQL命中: {'✓' if sql_hit else '✗'}")
        print(f"  耗时: {elapsed:.1f}s")

    # ═══ 汇总统计 ═══
    print("\n" + "=" * 70)
    print("  汇总统计")
    print("=" * 70)

    type_stats = {}
    for r in results:
        t = r["type"]
        if t not in type_stats:
            type_stats[t] = {"count": 0, "keyword_scores": [], "source_count": 0,
                             "no_hallucination_count": 0, "sql_hit_count": 0, "times": []}
        type_stats[t]["count"] += 1
        type_stats[t]["keyword_scores"].append(r["score"]["has_keyword"])
        type_stats[t]["source_count"] += r["score"]["has_source"]
        type_stats[t]["no_hallucination_count"] += r["score"]["no_hallucination"]
        type_stats[t]["sql_hit_count"] += 1 if r["sql_hit"] else 0
        type_stats[t]["times"].append(r["time"])

    print(f"\n{'类型':<12} {'数量':<6} {'关键词':<10} {'来源标注':<10} {'无幻觉':<8} {'SQL命中':<8} {'平均耗时':<10}")
    print("-" * 70)
    for t, s in type_stats.items():
        avg_kw = sum(s["keyword_scores"]) / len(s["keyword_scores"])
        source_rate = s["source_count"] / s["count"]
        hallu_rate = s["no_hallucination_count"] / s["count"]
        sql_rate = s["sql_hit_count"] / s["count"]
        avg_t = sum(s["times"]) / len(s["times"])
        print(f"{t:<12} {s['count']:<6} {avg_kw:.0%}{'':<6} {source_rate:.0%}{'':<6} "
              f"{hallu_rate:.0%}{'':<4} {sql_rate:.0%}{'':<4} {avg_t:.1f}s")

    total = len(results)
    avg_keyword = sum(r["score"]["has_keyword"] for r in results) / total
    total_source = sum(r["score"]["has_source"] for r in results)
    total_hallu = sum(r["score"]["no_hallucination"] for r in results)
    avg_time = total_time / total if total > 0 else 0

    total_stats = {
        "total": total,
        "avg_keyword": avg_keyword,
        "source_rate": total_source / total,
        "hallu_rate": total_hallu / total,
        "sql_rate": sql_hit_cases / total,
        "avg_time": avg_time,
        "api_ok": api_ok,
    }

    print("-" * 70)
    print(f"{'总体':<12} {total:<6} {avg_keyword:.0%}{'':<6} {total_source/total:.0%}{'':<6} "
          f"{total_hallu/total:.0%}{'':<4} {sql_hit_cases/total:.0%}{'':<4} {avg_time:.1f}s")
    print(f"\n  SQL 命中: {sql_hit_cases}/{total} ({sql_hit_cases/total:.0%})")
    print(f"  API 可用: {api_ok}/{total}")
    print(f"  平均响应时间: {avg_time:.1f}s")

    # ═══ SQL 覆盖分析 ═══
    print(f"\n{'─'*70}")
    print("  SQL 覆盖分析")
    print(f"{'─'*70}")
    sql_expected = [r for r in results if r.get("expect_sql")]
    sql_missed = [r for r in sql_expected if not r["sql_hit"]]
    print(f"  期望 SQL 命中: {len(sql_expected)} 题")
    print(f"  实际 SQL 命中: {len(sql_expected) - len(sql_missed)} 题")
    if sql_missed:
        print(f"  SQL 未命中 ({len(sql_missed)} 题):")
        for r in sql_missed:
            print(f"    [{r['id']}] {r['type']}: {r['question']}")

    # ═══ 保存 JSON ═══
    output_dir = os.path.join(os.path.dirname(__file__), "..", "..", "data")
    os.makedirs(output_dir, exist_ok=True)
    json_path = os.path.join(output_dir, "eval_results.json")

    save_results = []
    for r in results:
        save_results.append({
            "id": r["id"], "type": r["type"], "question": r["question"],
            "answer": r["answer_full"],
            "keywords_expected": r["keywords_expected"],
            "keywords_hit": r["keywords_hit"],
            "score": r["score"], "time": r["time"],
            "plan_log": {
                "intent": r["plan_log"].get("intent", ""),
                "graph_count": r["plan_log"].get("graph_count", 0),
                "sql_count": r["plan_log"].get("sql_count", 0),
                "vector_count": r["plan_log"].get("vector_count", 0),
                "sql_details": r["plan_log"].get("sql_details", []),
                "graph_details": r["plan_log"].get("graph_details", []),
                "entities_matched": r["plan_log"].get("entities_matched", []),
            } if r["plan_log"] else {},
            "expect_sql": r["expect_sql"], "sql_hit": r["sql_hit"],
        })

    with open(json_path, "w", encoding="utf-8") as f:
        json.dump(save_results, f, ensure_ascii=False, indent=2)
    print(f"\n[JSON] 已保存: {json_path}")

    # ═══ 生成报告（优先 DOCX，兜底 MD）═══
    docx_path = os.path.join(output_dir, "eval_report.docx")
    md_path = os.path.join(output_dir, "eval_report.md")

    if generate_docx(results, type_stats, total_stats, docx_path):
        print(f"[DOCX] 报告已保存: {docx_path}")
    else:
        print("[DOCX] python-docx 未安装，降级为 Markdown")
        generate_markdown(results, type_stats, total_stats, md_path)
        print(f"[MD] 报告已保存: {md_path}")


if __name__ == "__main__":
    main()
