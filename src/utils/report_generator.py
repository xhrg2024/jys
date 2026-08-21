"""
深度思考分析报告生成器：把一次深度思考的完整过程自动拼装为 Word 文档。
不额外调用 LLM——所有内容均来自深度思考流程已产生的数据（RACE 笔记、
工具调用结果、最终答案、来源索引），仅做结构化排版。
"""
import json
import re
from datetime import datetime
from pathlib import Path
from uuid import uuid4

REPORT_DIR = Path(__file__).resolve().parents[2] / "reports"
REPORT_DIR.mkdir(exist_ok=True)


def _extract_race_line(thinking, prefix):
    """从 RACE 笔记中提取某一行（如 R / A / C / E），兼容中文/英文冒号。"""
    for line in (thinking or "").split("\n"):
        s = line.strip()
        if s.startswith(prefix + "：") or s.startswith(prefix + ":"):
            return s[len(prefix) + 1:].strip()
    return ""


def _ref_label(entry):
    """尾注条目 → 一行可读文本，优先 ref_label，其次 desc，最后 label/编号兜底。"""
    if isinstance(entry, dict):
        return (
            entry.get("ref_label")
            or entry.get("desc")
            or entry.get("label")
            or None
        )
    return entry


def build_report_docx(question, thinking, tool_results, answer, source_index):
    """生成深度思考分析报告，返回保存后的 Path（失败抛异常，由调用方兜底）。

    结构：标题 / 一、问题理解 / 二、RACE 分析框架 / 三、工具调用过程 /
    四、最终分析结果 / 五、参考资料（尾注）。
    """
    from docx import Document

    doc = Document()

    # ── 标题 ──
    doc.add_heading("辑佚史深度分析报告", level=0)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    # ── 一、问题理解 ──
    doc.add_heading("一、问题理解", level=1)
    doc.add_paragraph(question or "（无）")
    r_line = _extract_race_line(thinking, "R")
    if r_line:
        doc.add_paragraph(f"问题类型与关键要素：{r_line}")

    # ── 二、RACE 分析框架 ──
    doc.add_heading("二、RACE 分析框架", level=1)
    doc.add_paragraph((thinking or "（无）").strip())

    # ── 三、工具调用过程 ──
    doc.add_heading("三、工具调用过程", level=1)
    if tool_results:
        doc.add_paragraph(f"本次深度思考共进行 {len(tool_results)} 次工具调用。")
        for i, tr in enumerate(tool_results, 1):
            name = tr.get("name", "?")
            args = tr.get("args", {})
            result = tr.get("result", "")
            round_no = tr.get("round", "?")
            doc.add_heading(f"3.{i} 第 {round_no} 轮 · {name}", level=2)
            if args:
                try:
                    args_text = json.dumps(args, ensure_ascii=False)
                except Exception:
                    args_text = str(args)
                doc.add_paragraph(f"参数：{args_text}")
            doc.add_paragraph("结果：")
            doc.add_paragraph((result or "（无）").strip())
    else:
        doc.add_paragraph("（本次未调用工具）")

    # ── 四、最终分析结果 ──
    doc.add_heading("四、最终分析结果", level=1)
    doc.add_paragraph((answer or "（无）").strip())

    # ── 五、参考资料（尾注）──
    doc.add_heading("五、参考资料", level=1)
    if source_index:
        def _sort_key(k):
            return int(k) if str(k).isdigit() else 0
        for num in sorted(source_index, key=_sort_key):
            label = _ref_label(source_index[num]) or f"来源 {num}"
            doc.add_paragraph(f"[{num}] {label}")
    else:
        doc.add_paragraph("（无）")

    # ── 保存 ──
    filename = f"{_new_id()}.docx"
    out_path = REPORT_DIR / filename
    doc.save(str(out_path))
    return out_path


def _new_id():
    return f"{datetime.now().strftime('%Y%m%d_%H%M%S')}_{uuid4().hex[:8]}"


def save_session_data(question, thinking, tool_results, answer, source_index):
    """保存一次深度思考的完整数据（供用户稍后手动生成报告）。返回 session_id。"""
    session_id = _new_id()
    data = {
        "question": question,
        "thinking": thinking,
        "tool_results": tool_results,
        "answer": answer,
        "source_index": source_index,
    }
    path = REPORT_DIR / f"{session_id}.json"
    with open(path, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False)
    return session_id


def load_session_data(session_id):
    """读取已保存的会话数据；不存在则抛 FileNotFoundError。"""
    path = REPORT_DIR / f"{session_id}.json"
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def generate_report_from_session(session_id):
    """从已保存的会话数据生成 Word 报告，返回 docx 的 Path。"""
    data = load_session_data(session_id)
    return build_report_docx(
        question=data.get("question"),
        thinking=data.get("thinking"),
        tool_results=data.get("tool_results") or [],
        answer=data.get("answer"),
        source_index=data.get("source_index") or {},
    )
